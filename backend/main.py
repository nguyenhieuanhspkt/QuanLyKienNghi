import difflib
import io
import json
from datetime import date
from pathlib import Path
from typing import List, Optional

from docx import Document as DocxDocument
from fastapi import FastAPI, HTTPException, UploadFile, File, Form
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from fastapi.staticfiles import StaticFiles
from openpyxl import Workbook, load_workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter
from pydantic import BaseModel

app = FastAPI(title="QuanLyKienNghi API", version="1.0.0")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:5173", "http://localhost:3000"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

DATA_FILE = Path(__file__).parent / "data.json"

DON_VI_LIST = ["PXVH", "PXNL", "PKT", "Hoàng Anh", "Thanh Tuấn"]
TRANG_THAI_LIST = [
    "phong_ktat_len_phieu",
    "da_sua_chua_theo_doi",
    "da_cap_vat_tu",
    "cho_vat_tu",
]
THOI_HAN_VT_LIST = [
    "duoi_1_thang",
    "2_3_thang",
    "3_6_thang",
    "tren_6_thang",
    "chua_co_han",
]


def doc_data() -> dict:
    with open(DATA_FILE, "r", encoding="utf-8") as f:
        return json.load(f)


def ghi_data(data: dict):
    with open(DATA_FILE, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)


class KienNghiCreate(BaseModel):
    don_vi: str
    noi_dung: str
    tuan_truoc: Optional[str] = ""
    tuan_nay: Optional[str] = ""
    trang_thai: str
    thoi_han_vt: Optional[str] = None
    tuan_phat_sinh: str


class KienNghiUpdate(BaseModel):
    don_vi: Optional[str] = None
    noi_dung: Optional[str] = None
    tuan_truoc: Optional[str] = None
    tuan_nay: Optional[str] = None
    trang_thai: Optional[str] = None
    thoi_han_vt: Optional[str] = None
    tuan_phat_sinh: Optional[str] = None
    hoan_thanh: Optional[bool] = None
    ngay_hoan_thanh: Optional[str] = None


# ─── CRUD ────────────────────────────────────────────────────────────────────

@app.get("/api/v1/kien-nghi", summary="Lấy danh sách kiến nghị")
def get_all(hoan_thanh: Optional[bool] = None):
    data = doc_data()
    items = data["kien_nghi"]
    if hoan_thanh is not None:
        items = [kn for kn in items if kn["hoan_thanh"] == hoan_thanh]
    return items


@app.get("/api/v1/kien-nghi/{kn_id}", summary="Lấy 1 kiến nghị")
def get_one(kn_id: str):
    data = doc_data()
    for kn in data["kien_nghi"]:
        if kn["id"] == kn_id:
            return kn
    raise HTTPException(404, "Không tìm thấy kiến nghị")


@app.post("/api/v1/kien-nghi", status_code=201, summary="Thêm kiến nghị mới")
def create(body: KienNghiCreate):
    data = doc_data()
    # Backward compat: khởi tạo id_counter từ max ID hiện có nếu thiếu
    if "id_counter" not in data:
        max_num = 0
        for kn in data["kien_nghi"]:
            try:
                max_num = max(max_num, int(kn["id"].split("-")[2]))
            except (IndexError, ValueError):
                pass
        data["id_counter"] = max_num
    data["id_counter"] += 1
    so_thu_tu = data["id_counter"]
    nam = date.today().year
    kn = {
        "id": f"KN-{nam}-{so_thu_tu:03d}",
        "don_vi": body.don_vi,
        "noi_dung": body.noi_dung,
        "tuan_truoc": body.tuan_truoc or "",
        "tuan_nay": body.tuan_nay or "",
        "trang_thai": body.trang_thai,
        "thoi_han_vt": body.thoi_han_vt,
        "tuan_phat_sinh": body.tuan_phat_sinh,
        "hoan_thanh": False,
        "ngay_hoan_thanh": None,
        "lich_su": [],
    }
    data["kien_nghi"].append(kn)
    ghi_data(data)
    return kn


@app.put("/api/v1/kien-nghi/{kn_id}", summary="Cập nhật kiến nghị")
def update(kn_id: str, body: KienNghiUpdate):
    data = doc_data()
    for idx, kn in enumerate(data["kien_nghi"]):
        if kn["id"] == kn_id:
            update_fields = body.model_dump(exclude_unset=True)
            # Lưu lịch sử trước khi cập nhật
            if any(k in update_fields for k in ("tuan_nay", "trang_thai", "hoan_thanh")):
                snapshot = {k: kn[k] for k in ("tuan_nay", "trang_thai", "hoan_thanh")}
                kn["lich_su"].append(snapshot)
            kn.update(update_fields)
            data["kien_nghi"][idx] = kn
            ghi_data(data)
            return kn
    raise HTTPException(404, "Không tìm thấy kiến nghị")


@app.delete("/api/v1/kien-nghi/{kn_id}", status_code=204, summary="Xóa kiến nghị")
def delete(kn_id: str):
    data = doc_data()
    before = len(data["kien_nghi"])
    data["kien_nghi"] = [kn for kn in data["kien_nghi"] if kn["id"] != kn_id]
    if len(data["kien_nghi"]) == before:
        raise HTTPException(404, "Không tìm thấy kiến nghị")
    ghi_data(data)


# ─── Cập nhật tuần mới ───────────────────────────────────────────────────────

@app.post("/api/v1/cap-nhat-tuan-moi", summary="Kế thừa tuan_nay → tuan_truoc cho tuần tiếp")
def cap_nhat_tuan_moi():
    data = doc_data()
    count = 0
    for kn in data["kien_nghi"]:
        if not kn["hoan_thanh"]:
            kn["tuan_truoc"] = kn["tuan_nay"]
            kn["tuan_nay"] = ""
            count += 1
    ghi_data(data)
    return {"message": f"Đã cập nhật {count} kiến nghị chưa hoàn thành"}


# ─── Xuất Excel ──────────────────────────────────────────────────────────────

LABEL_TRANG_THAI = {
    "phong_ktat_len_phieu": "Phòng KTAT đang lên phiếu YC",
    "da_sua_chua_theo_doi": "Đã sửa chữa, đang theo dõi",
    "da_cap_vat_tu": "Đã cấp vật tư",
    "cho_vat_tu": "Đợi cấp vật tư",
}
LABEL_THOI_HAN = {
    "duoi_1_thang": "< 1 tháng",
    "2_3_thang": "2–3 tháng",
    "3_6_thang": "3–6 tháng",
    "tren_6_thang": "> 6 tháng",
    "chua_co_han": "Chưa có hạn",
}

def _cell_border():
    s = Side(style="thin")
    return Border(left=s, right=s, top=s, bottom=s)

def _header_fill():
    return PatternFill("solid", fgColor="1E3A5F")

def _set_col_widths(ws, widths):
    for i, w in enumerate(widths, 1):
        ws.column_dimensions[get_column_letter(i)].width = w

def _write_header_row(ws, headers):
    border = _cell_border()
    fill = _header_fill()
    font = Font(bold=True, color="FFFFFF", name="Arial", size=10)
    align = Alignment(horizontal="center", vertical="center", wrap_text=True)
    for col, h in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col, value=h)
        cell.font = font
        cell.fill = fill
        cell.alignment = align
        cell.border = border
    ws.row_dimensions[1].height = 30

def _wrap_cell(cell, value):
    cell.value = value
    cell.alignment = Alignment(wrap_text=True, vertical="top")
    cell.border = _cell_border()
    cell.font = Font(name="Arial", size=10)

def _build_sheet_chua_ht(ws, items):
    headers = ["STT", "Mã KN", "Đơn vị", "Nội dung kiến nghị",
               "Xử lý tuần trước", "Cập nhật tuần này",
               "Trạng thái", "Tuần phát sinh"]
    _write_header_row(ws, headers)
    _set_col_widths(ws, [5, 13, 10, 40, 30, 30, 28, 14])

    alt_fill = PatternFill("solid", fgColor="EFF6FF")
    for i, kn in enumerate(items, 1):
        row = i + 1
        tt = LABEL_TRANG_THAI.get(kn["trang_thai"], kn["trang_thai"])
        if kn["trang_thai"] == "cho_vat_tu" and kn.get("thoi_han_vt"):
            tt += f"\n({LABEL_THOI_HAN.get(kn['thoi_han_vt'], kn['thoi_han_vt'])})"

        values = [i, kn["id"], kn["don_vi"], kn["noi_dung"],
                  kn.get("tuan_truoc", ""), kn.get("tuan_nay", ""),
                  tt, kn["tuan_phat_sinh"]]
        for col, val in enumerate(values, 1):
            cell = ws.cell(row=row, column=col)
            _wrap_cell(cell, val)
            if i % 2 == 0:
                cell.fill = alt_fill
        ws.row_dimensions[row].height = 55

    # Thống kê theo nhóm
    from collections import Counter
    nhom = Counter(kn["trang_thai"] for kn in items)
    stat_row = len(items) + 3
    ws.cell(row=stat_row, column=1, value="THỐNG KÊ THEO TRẠNG THÁI").font = Font(bold=True, name="Arial", size=10)
    for j, (key, label) in enumerate(LABEL_TRANG_THAI.items(), 1):
        r = stat_row + j
        ws.cell(row=r, column=1, value=label).font = Font(name="Arial", size=10)
        ws.cell(row=r, column=2, value=nhom.get(key, 0)).font = Font(bold=True, name="Arial", size=10)
    ws.cell(row=stat_row + len(LABEL_TRANG_THAI) + 1, column=1, value="Tổng cộng").font = Font(bold=True, name="Arial", size=10)
    ws.cell(row=stat_row + len(LABEL_TRANG_THAI) + 1, column=2, value=len(items)).font = Font(bold=True, name="Arial", size=10)

def _build_sheet_da_ht(ws, items):
    headers = ["STT", "Mã KN", "Đơn vị", "Nội dung kiến nghị",
               "Xử lý cuối", "Tuần phát sinh", "Ngày hoàn thành"]
    _write_header_row(ws, headers)
    _set_col_widths(ws, [5, 13, 10, 45, 35, 14, 14])

    alt_fill = PatternFill("solid", fgColor="F0FDF4")
    for i, kn in enumerate(items, 1):
        row = i + 1
        values = [i, kn["id"], kn["don_vi"], kn["noi_dung"],
                  kn.get("tuan_nay", ""), kn["tuan_phat_sinh"],
                  kn.get("ngay_hoan_thanh", "")]
        for col, val in enumerate(values, 1):
            cell = ws.cell(row=row, column=col)
            _wrap_cell(cell, val)
            if i % 2 == 0:
                cell.fill = alt_fill
        ws.row_dimensions[row].height = 55


@app.get("/api/v1/xuat-excel", summary="Xuất báo cáo Excel 2 danh sách")
def xuat_excel():
    data = doc_data()
    items = data["kien_nghi"]
    chua_ht = [kn for kn in items if not kn["hoan_thanh"]]
    da_ht = [kn for kn in items if kn["hoan_thanh"]]

    today = date.today().strftime("%d/%m/%Y")
    wb = Workbook()

    ws1 = wb.active
    ws1.title = "Chưa hoàn thành"
    ws1.sheet_properties.tabColor = "2563EB"
    _build_sheet_chua_ht(ws1, chua_ht)

    ws2 = wb.create_sheet("Đã hoàn thành")
    ws2.sheet_properties.tabColor = "16A34A"
    _build_sheet_da_ht(ws2, da_ht)

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)

    ten_file = f"BaoCaoKienNghi_{date.today().strftime('%Y%m%d')}.xlsx"
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f"attachment; filename*=UTF-8''{ten_file}"},
    )


# ─── Import / Preview Excel ──────────────────────────────────────────────────

_LABEL_TO_TRANG_THAI = {v: k for k, v in LABEL_TRANG_THAI.items()}
_LABEL_TO_THOI_HAN = {v: k for k, v in LABEL_THOI_HAN.items()}
# Alias cho key không chuẩn từ hệ thống khác
_TRANG_THAI_ALIAS = {
    "dang_theo_doi": "da_sua_chua_theo_doi",
    "theo_doi": "da_sua_chua_theo_doi",
}


def _parse_trang_thai(cell_value):
    """Chuyển nhãn / key trạng thái → (trang_thai_key, thoi_han_vt_key)."""
    if not cell_value:
        return "cho_vat_tu", None
    s = str(cell_value).strip()
    if s in TRANG_THAI_LIST:
        return s, None
    if s in _TRANG_THAI_ALIAS:
        return _TRANG_THAI_ALIAS[s], None
    thoi_han = None
    if "\n" in s:
        label, rest = s.split("\n", 1)
        thoi_han = _LABEL_TO_THOI_HAN.get(rest.strip().strip("()"))
    else:
        label = s
    return _LABEL_TO_TRANG_THAI.get(label.strip(), "cho_vat_tu"), thoi_han


def _cell_str(row_values, col_1based: int) -> str:
    """Lấy giá trị ô (1-based index) từ tuple values_only, trả về str."""
    if not col_1based or col_1based < 1 or col_1based > len(row_values):
        return ""
    val = row_values[col_1based - 1]
    return str(val).strip() if val is not None else ""


def _ensure_id_counter(data: dict):
    if "id_counter" not in data:
        max_num = 0
        for kn in data["kien_nghi"]:
            try:
                max_num = max(max_num, int(kn["id"].split("-")[2]))
            except (IndexError, ValueError):
                pass
        data["id_counter"] = max_num


@app.post("/api/v1/preview-excel", summary="Xem trước file Excel trước khi import")
async def preview_excel(file: UploadFile = File(...)):
    if not file.filename.endswith(".xlsx"):
        raise HTTPException(400, "Chỉ hỗ trợ file .xlsx")
    content = await file.read()
    try:
        wb = load_workbook(io.BytesIO(content), data_only=True)
    except Exception:
        raise HTTPException(400, "Không đọc được file Excel")

    result = []
    for ws in wb.worksheets:
        rows = []
        for i, row in enumerate(ws.iter_rows(min_row=1, values_only=True)):
            if i >= 7:
                break
            rows.append([str(c).strip() if c is not None else "" for c in row])
        result.append({"name": ws.title, "rows": rows})
    return result


@app.post("/api/v1/import-excel", summary="Import kiến nghị từ file Excel (.xlsx)")
async def import_excel(
    file: UploadFile = File(...),
    mapping: str = Form(default=None),
):
    if not file.filename.endswith(".xlsx"):
        raise HTTPException(400, "Chỉ hỗ trợ file .xlsx")
    content = await file.read()
    try:
        wb = load_workbook(io.BytesIO(content), data_only=True)
    except Exception:
        raise HTTPException(400, "Không đọc được file Excel")

    data = doc_data()
    _ensure_id_counter(data)
    existing_ids = {kn["id"] for kn in data["kien_nghi"]}
    imported = skipped_dup = skipped_empty = 0

    if mapping:
        mp = json.loads(mapping)
        sheet_idx = int(mp.get("sheet_index", 0))
        header_row = int(mp.get("header_row", 1))
        hoan_thanh = bool(mp.get("hoan_thanh", False))
        col_ma_kn = mp.get("col_ma_kn")
        col_don_vi = mp.get("col_don_vi")
        don_vi_fixed = mp.get("don_vi_fixed")
        col_noi_dung = mp.get("col_noi_dung")
        col_tuan_truoc = mp.get("col_tuan_truoc")
        col_tuan_nay = mp.get("col_tuan_nay")
        col_trang_thai = mp.get("col_trang_thai")
        trang_thai_fixed = mp.get("trang_thai_fixed", "cho_vat_tu")
        thoi_han_vt_fixed = mp.get("thoi_han_vt_fixed")
        col_tuan_phat_sinh = mp.get("col_tuan_phat_sinh")
        col_ngay_hoan_thanh = mp.get("col_ngay_hoan_thanh")

        ws = wb.worksheets[sheet_idx] if sheet_idx < len(wb.worksheets) else wb.worksheets[0]

        for row_num, row_vals in enumerate(ws.iter_rows(min_row=1, values_only=True), start=1):
            if row_num <= header_row:
                continue
            noi_dung = _cell_str(row_vals, col_noi_dung)
            if not noi_dung:
                skipped_empty += 1
                continue

            ma_kn = _cell_str(row_vals, col_ma_kn) if col_ma_kn else ""
            if not ma_kn:
                data["id_counter"] += 1
                ma_kn = f"KN-{date.today().year}-{data['id_counter']:03d}"
            if ma_kn in existing_ids:
                skipped_dup += 1
                continue

            don_vi = don_vi_fixed or _cell_str(row_vals, col_don_vi) or ""

            if col_trang_thai:
                trang_thai, thoi_han = _parse_trang_thai(_cell_str(row_vals, col_trang_thai))
            else:
                trang_thai = trang_thai_fixed or "cho_vat_tu"
                thoi_han = thoi_han_vt_fixed if trang_thai == "cho_vat_tu" else None

            kn = {
                "id": ma_kn,
                "don_vi": don_vi,
                "noi_dung": noi_dung,
                "tuan_truoc": _cell_str(row_vals, col_tuan_truoc),
                "tuan_nay": _cell_str(row_vals, col_tuan_nay),
                "trang_thai": trang_thai,
                "thoi_han_vt": thoi_han or None,
                "tuan_phat_sinh": _cell_str(row_vals, col_tuan_phat_sinh),
                "hoan_thanh": hoan_thanh,
                "ngay_hoan_thanh": _cell_str(row_vals, col_ngay_hoan_thanh) if hoan_thanh else None,
                "lich_su": [],
            }
            data["kien_nghi"].append(kn)
            existing_ids.add(ma_kn)
            imported += 1
    else:
        # Fallback: format cố định (format xuất của app)
        for sheet_idx, hoan_thanh in [(0, False), (1, True)]:
            if sheet_idx >= len(wb.worksheets):
                break
            ws = wb.worksheets[sheet_idx]
            for row_vals in ws.iter_rows(min_row=2, values_only=True):
                ma_kn = _cell_str(row_vals, 2)
                if not ma_kn or not ma_kn.startswith("KN-"):
                    skipped_empty += 1
                    continue
                if ma_kn in existing_ids:
                    skipped_dup += 1
                    continue
                if not hoan_thanh:
                    trang_thai, thoi_han = _parse_trang_thai(_cell_str(row_vals, 7))
                    kn = {
                        "id": ma_kn, "don_vi": _cell_str(row_vals, 3),
                        "noi_dung": _cell_str(row_vals, 4), "tuan_truoc": _cell_str(row_vals, 5),
                        "tuan_nay": _cell_str(row_vals, 6), "trang_thai": trang_thai,
                        "thoi_han_vt": thoi_han, "tuan_phat_sinh": _cell_str(row_vals, 8),
                        "hoan_thanh": False, "ngay_hoan_thanh": None, "lich_su": [],
                    }
                else:
                    kn = {
                        "id": ma_kn, "don_vi": _cell_str(row_vals, 3),
                        "noi_dung": _cell_str(row_vals, 4), "tuan_truoc": "",
                        "tuan_nay": _cell_str(row_vals, 5), "trang_thai": "da_cap_vat_tu",
                        "thoi_han_vt": None, "tuan_phat_sinh": _cell_str(row_vals, 6),
                        "hoan_thanh": True, "ngay_hoan_thanh": _cell_str(row_vals, 7), "lich_su": [],
                    }
                data["kien_nghi"].append(kn)
                existing_ids.add(ma_kn)
                imported += 1

    # Sync id_counter với max ID thực tế
    for kn in data["kien_nghi"]:
        try:
            data["id_counter"] = max(data["id_counter"], int(kn["id"].split("-")[2]))
        except (IndexError, ValueError):
            pass

    ghi_data(data)
    return {
        "imported": imported,
        "skipped_dup": skipped_dup,
        "skipped_empty": skipped_empty,
        "message": f"Import thành công {imported} kiến nghị, bỏ qua {skipped_dup} trùng ID.",
    }


# ─── Import Word ─────────────────────────────────────────────────────────────

_DON_VI_WORD_MAP = {
    "Phân xưởng Vận hành": "PXVH",
    "Phân xưởng nhiên liệu": "PXNL",
    "Phòng Kỹ thuật và An toàn": "PKT",
}


def _doc_cell_text(cell) -> str:
    return "\n".join(p.text.strip() for p in cell.paragraphs if p.text.strip())


def _extract_ket_qua_khvt(col5_text: str) -> str:
    for line in col5_text.split("\n"):
        s = line.strip()
        if s.startswith("→") and "P.KHVT" in s and ":" in s:
            idx = s.find(":", s.find("P.KHVT"))
            if idx != -1:
                return s[idx + 1:].strip()
    return ""


def _fuzzy_score(a: str, b: str) -> float:
    return difflib.SequenceMatcher(None, a.lower(), b.lower()).ratio()


@app.post("/api/v1/import-word", summary="Phân tích file Word KHVT, fuzzy match với DB")
async def import_word_preview(file: UploadFile = File(...)):
    if not file.filename.endswith(".docx"):
        raise HTTPException(400, "Chỉ hỗ trợ file .docx")
    content = await file.read()
    try:
        doc = DocxDocument(io.BytesIO(content))
    except Exception:
        raise HTTPException(400, "Không đọc được file Word")

    if len(doc.tables) < 3:
        raise HTTPException(400, "File Word cần có ít nhất 3 bảng (Bảng 3 là nguồn dữ liệu tồn tại đơn vị)")

    table = doc.tables[2]
    current_don_vi = None
    word_items = []

    for row in table.rows:
        cells = row.cells
        if len(cells) < 5:
            continue
        stt_raw = _doc_cell_text(cells[0]).strip()
        noi_dung = _doc_cell_text(cells[1]).strip()
        de_xuat = _doc_cell_text(cells[4]).strip()

        if not stt_raw:
            continue

        try:
            int(stt_raw)
            is_numbered = True
        except ValueError:
            is_numbered = False

        if not is_numbered:
            # Exact case-insensitive match — tránh nhầm "Phân xưởng Vận hành" với
            # "Phân xưởng vận hành điện Mặt trời..." hay các đơn vị khác
            current_don_vi = None
            for name, code in _DON_VI_WORD_MAP.items():
                if stt_raw.lower().strip() == name.lower().strip():
                    current_don_vi = code
                    break
            continue

        if noi_dung and current_don_vi:
            word_items.append({
                "don_vi": current_don_vi,
                "noi_dung_word": noi_dung,
                "tuan_nay_moi": de_xuat,
                "ket_qua_khvt": _extract_ket_qua_khvt(de_xuat),
            })

    if not word_items:
        raise HTTPException(
            422,
            "Không đọc được dữ liệu từ Bảng 2. Kiểm tra lại file Word: cần có hàng tên đơn vị "
            "(PXVH/PXNL/PKT) và hàng dữ liệu với STT là số nguyên.",
        )

    data = doc_data()
    chua_ht = [kn for kn in data["kien_nghi"] if not kn["hoan_thanh"]]
    chac_chan, khong_chac, khong_khop = [], [], []

    for idx, item in enumerate(word_items):
        scored = sorted(
            [(_fuzzy_score(item["noi_dung_word"], kn["noi_dung"]), kn) for kn in chua_ht],
            key=lambda x: -x[0],
        )
        base = {
            "word_index": idx,
            "don_vi": item["don_vi"],
            "noi_dung_word": item["noi_dung_word"],
            "tuan_nay_moi": item["tuan_nay_moi"],
            "ket_qua_khvt": item["ket_qua_khvt"],
        }
        best_score = scored[0][0] if scored else 0

        if best_score >= 0.75:
            kn = scored[0][1]
            chac_chan.append({
                **base,
                "match": {
                    "kn_id": kn["id"],
                    "noi_dung": kn["noi_dung"],
                    "tuan_nay_hien_tai": kn.get("tuan_nay", ""),
                    "score": round(best_score, 3),
                },
            })
        elif best_score >= 0.45:
            top3 = [
                {"kn_id": kn["id"], "noi_dung": kn["noi_dung"], "score": round(sc, 3)}
                for sc, kn in scored[:3]
                if sc >= 0.45
            ]
            khong_chac.append({**base, "top3_matches": top3})
        else:
            khong_khop.append(base)

    return {"chac_chan": chac_chan, "khong_chac": khong_chac, "khong_khop": khong_khop}


class WordActionItem(BaseModel):
    word_index: int
    action: str
    kn_id: Optional[str] = None
    don_vi: str
    noi_dung_word: str
    tuan_nay_moi: str
    ket_qua_khvt: str


class XacNhanWordRequest(BaseModel):
    actions: List[WordActionItem]


@app.post("/api/v1/import-word/xac-nhan", summary="Xác nhận import Word vào DB")
def import_word_xac_nhan(body: XacNhanWordRequest):
    data = doc_data()
    _ensure_id_counter(data)
    merged = created = 0

    for act in body.actions:
        if act.action == "merge":
            if not act.kn_id:
                continue
            for kn in data["kien_nghi"]:
                if kn["id"] == act.kn_id:
                    snapshot = {k: kn.get(k) for k in ("tuan_nay", "trang_thai", "hoan_thanh", "ngay_hoan_thanh")}
                    kn.setdefault("lich_su", []).append(snapshot)
                    kn["tuan_nay"] = act.tuan_nay_moi
                    if act.ket_qua_khvt:
                        kn["ket_qua_khvt"] = act.ket_qua_khvt
                    # Nếu KN đã hoàn thành → kích hoạt lại vào vòng đời
                    if kn.get("hoan_thanh"):
                        kn["hoan_thanh"] = False
                        kn["ngay_hoan_thanh"] = None
                        kn["trang_thai"] = "phong_ktat_len_phieu"
                    merged += 1
                    break
        elif act.action == "tao_moi":
            data["id_counter"] += 1
            nam = date.today().year
            iso_week = date.today().isocalendar()[1]
            kn_new = {
                "id": f"KN-{nam}-{data['id_counter']:03d}",
                "don_vi": act.don_vi,
                "noi_dung": act.noi_dung_word,
                "tuan_truoc": "",
                "tuan_nay": act.tuan_nay_moi,
                "trang_thai": "phong_ktat_len_phieu",
                "thoi_han_vt": None,
                "tuan_phat_sinh": f"Tuần {iso_week}/{nam}",
                "hoan_thanh": False,
                "ngay_hoan_thanh": None,
                "ket_qua_ktat": "",
                "ket_qua_khvt": act.ket_qua_khvt,
                "so_pyc": None,
                "lich_su": [],
            }
            data["kien_nghi"].append(kn_new)
            created += 1

    ghi_data(data)
    return {
        "merged": merged,
        "created": created,
        "message": f"Đã cập nhật {merged} KN, tạo mới {created} KN.",
    }


# ─── Bảng chia việc ──────────────────────────────────────────────────────────

class BangChiaViecItem(BaseModel):
    id: str
    don_vi_phan_cong: Optional[str] = "CA_HAI"

class XuatBangChiaViecRequest(BaseModel):
    xuat: List[BangChiaViecItem]
    hoan_thanh: List[str] = []
    xoa: List[str] = []


@app.post("/api/v1/xuat-bang-chia-viec", summary="Duyệt + xuất bảng chia việc Excel")
def xuat_bang_chia_viec(body: XuatBangChiaViecRequest):
    data = doc_data()
    today_str = date.today().strftime("%d/%m/%Y")

    # 1. Xóa KN khỏi DB
    xoa_set = set(body.xoa)
    data["kien_nghi"] = [kn for kn in data["kien_nghi"] if kn["id"] not in xoa_set]

    # 2. Đánh dấu hoàn thành
    ht_set = set(body.hoan_thanh)
    for kn in data["kien_nghi"]:
        if kn["id"] in ht_set:
            kn["hoan_thanh"] = True
            kn["ngay_hoan_thanh"] = today_str

    # 3. Lưu don_vi_phan_cong vào DB, thu thập danh sách xuất (giữ thứ tự từ request)
    xuat_order = {item.id: (i, item.don_vi_phan_cong) for i, item in enumerate(body.xuat)}
    for kn in data["kien_nghi"]:
        if kn["id"] in xuat_order:
            kn["don_vi_phan_cong"] = xuat_order[kn["id"]][1]

    kn_xuat = sorted(
        [kn for kn in data["kien_nghi"] if kn["id"] in xuat_order],
        key=lambda kn: xuat_order[kn["id"]][0],
    )

    ghi_data(data)

    # 4. Tạo Excel
    wb = Workbook()
    ws = wb.active
    ws.title = "Bảng chia việc"

    headers = [
        "STT", "Mã KN", "Đơn vị", "Nội dung kiến nghị", "Kết quả tuần trước",
        "Tình hình KTAT xử lý", "Đã có PYC?", "Số PYC",
        "Thời hạn cấp VT", "Tình trạng cấp VT", "Ghi chú KHVT",
    ]
    _write_header_row(ws, headers)
    _set_col_widths(ws, [5, 13, 10, 40, 30, 30, 12, 10, 20, 25, 25])

    yellow_fill = PatternFill("solid", fgColor="FFD966")
    blue_fill   = PatternFill("solid", fgColor="9FC5E8")
    grey_fill   = PatternFill("solid", fgColor="D9D9D9")
    border = _cell_border()

    for i, kn in enumerate(kn_xuat, 1):
        row = i + 1
        phan_cong = xuat_order[kn["id"]][1] or "CA_HAI"

        for col, val in enumerate(
            [i, kn["id"], kn["don_vi"], kn["noi_dung"], kn.get("tuan_truoc", "")], 1
        ):
            _wrap_cell(ws.cell(row=row, column=col), val)

        ktat_fill = yellow_fill if phan_cong in ("KTAT", "CA_HAI") else grey_fill
        khvt_fill = blue_fill   if phan_cong in ("KHVT", "CA_HAI") else grey_fill

        for col in range(6, 9):   # F–H: vùng KTAT
            c = ws.cell(row=row, column=col, value="")
            c.fill = ktat_fill
            c.border = border
            c.alignment = Alignment(wrap_text=True, vertical="top")
            c.font = Font(name="Arial", size=10)

        for col in range(9, 12):  # I–K: vùng KHVT
            c = ws.cell(row=row, column=col, value="")
            c.fill = khvt_fill
            c.border = border
            c.alignment = Alignment(wrap_text=True, vertical="top")
            c.font = Font(name="Arial", size=10)

        ws.row_dimensions[row].height = 55

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)

    ten_file = f"BangChiaViec_{date.today().strftime('%Y%m%d')}.xlsx"
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f"attachment; filename*=UTF-8''{ten_file}"},
    )


# ─── Thống kê ────────────────────────────────────────────────────────────────

@app.get("/api/v1/thong-ke", summary="Thống kê theo trạng thái")
def thong_ke():
    data = doc_data()
    items = data["kien_nghi"]
    chua_ht = [kn for kn in items if not kn["hoan_thanh"]]
    da_ht = [kn for kn in items if kn["hoan_thanh"]]

    nhom = {}
    for kn in chua_ht:
        tt = kn["trang_thai"]
        nhom[tt] = nhom.get(tt, 0) + 1

    cho_vat_tu_nhom = {}
    for kn in chua_ht:
        if kn["trang_thai"] == "cho_vat_tu":
            th = kn.get("thoi_han_vt") or "chua_co_han"
            cho_vat_tu_nhom[th] = cho_vat_tu_nhom.get(th, 0) + 1

    return {
        "tong": len(items),
        "chua_hoan_thanh": len(chua_ht),
        "da_hoan_thanh": len(da_ht),
        "theo_trang_thai": nhom,
        "cho_vat_tu_theo_thoi_han": cho_vat_tu_nhom,
    }


# ─── Serve Frontend (production) ─────────────────────────────────────────────
# Đặt cuối cùng để không ghi đè các API route ở trên
_FRONTEND_DIST = Path(__file__).parent.parent / "frontend" / "dist"
if _FRONTEND_DIST.exists():
    app.mount("/", StaticFiles(directory=_FRONTEND_DIST, html=True), name="static")
